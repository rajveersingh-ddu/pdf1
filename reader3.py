import io
import re
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import pdfplumber
from docx import Document

app = Flask(__name__)
CORS(app)  # This allows your WordPress site to talk to this script

def clean_for_ats(text):
    # Remove non-ASCII (icons, symbols) and normalize whitespace
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    return re.sub(r'\s+', ' ', text).strip()

@app.route('/convert', methods=['POST'])
def convert():
    if 'resume' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['resume']
    
    # 1. Extract text from PDF
    extracted_text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            content = page.extract_text()
            if content:
                extracted_text += content + "\n"

    # 2. Sanitize text
    clean_text = clean_for_ats(extracted_text)

    # 3. Create a standard DOCX
    doc = Document()
    doc.add_heading('ATS Optimized Resume', 0)
    doc.add_paragraph(clean_text)

    # 4. Save to memory and return
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="ATS_Optimized.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == '__main__':
    app.run()
